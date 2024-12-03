import os
import random
from datetime import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docxcompose.composer import Composer

from configs.question.export.doc_export import (
    COVER_PAGE_STYLES,
    SUBJECT_NAME_HEADING_STYLES,
    PRACTICE_EXAMINATION_HEADING_STYLES,
    MCQ_HEADING_STYLES,
    MCQ_HEADING_ABBR_STYLES,
    ESSAY_QUESTION_HEADING_STYLES,
    RESTRICTED_RESPONSE_ESSAY_QUESTION_HEADING_STYLES,
    EXTENDED_RESPONSE_ESSAY_QUESTION_HEADING_STYLES,
    ANSWER_KEY_HEADING_STYLES,
    EXPORT_FILE_SAVE_DIR_PATH,
    COVER_IMAGE_METADATA
)
from utils.datetime import get_formatted_datetime


def is_style_available(document: Document, style_name: str) -> bool:
    for style in document.styles:
        if style.name == style_name:
            return True
    return False


def init_cover_page(document: Document, style_name: str = 'cover-page') -> None:
    if not is_style_available(document, style_name):
        cover_page_style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        cover_page_style.paragraph_format.alignment = COVER_PAGE_STYLES['paragraph']['alignment']


def init_topic_heading(document: Document, style_name: str = 'topic-heading') -> None:
    if not is_style_available(document, style_name):
        topic_heading_style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        topic_heading_style.paragraph_format.alignment = SUBJECT_NAME_HEADING_STYLES['paragraph']['alignment']
        topic_heading_style.font.name = SUBJECT_NAME_HEADING_STYLES['font']['name']
        topic_heading_style.font.size = SUBJECT_NAME_HEADING_STYLES['font']['size']
        topic_heading_style.font.color.rgb = SUBJECT_NAME_HEADING_STYLES['font']['color']


def init_mcq_heading(document: Document, style_name: str = 'mcq-heading') -> None:
    if not is_style_available(document, style_name):
        mcq_question_item_heading_style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        mcq_question_item_heading_abbr_style = document.styles.add_style(f'{style_name}-abbr', WD_STYLE_TYPE.CHARACTER)
        mcq_question_item_heading_style.paragraph_format.alignment = MCQ_HEADING_STYLES['paragraph']['alignment']
        mcq_question_item_heading_style.font.name = MCQ_HEADING_STYLES['font']['name']
        mcq_question_item_heading_abbr_style.font.name = MCQ_HEADING_ABBR_STYLES['font']['name']
        mcq_question_item_heading_style.font.size = MCQ_HEADING_STYLES['font']['size']
        mcq_question_item_heading_abbr_style.font.size = MCQ_HEADING_ABBR_STYLES['font']['size']
        mcq_question_item_heading_style.font.color.rgb = MCQ_HEADING_STYLES['font']['color']
        mcq_question_item_heading_abbr_style.font.color.rgb = MCQ_HEADING_ABBR_STYLES['font']['color']


def init_essay_heading(document: Document, style_name: str = 'essay-question-heading') -> None:
    if not is_style_available(document, style_name):
        essay_question_heading_style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        essay_question_heading_style.paragraph_format.alignment = ESSAY_QUESTION_HEADING_STYLES['paragraph'][
            'alignment']
        essay_question_heading_style.font.name = ESSAY_QUESTION_HEADING_STYLES['font']['name']
        essay_question_heading_style.font.size = ESSAY_QUESTION_HEADING_STYLES['font']['size']
        essay_question_heading_style.font.color.rgb = ESSAY_QUESTION_HEADING_STYLES['font']['color']


def init_restricted_essay_heading(document: Document, style_name: str = 'restricted-essay-question-heading') -> None:
    if not is_style_available(document, style_name):
        restricted_essay_question_heading_style = document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        restricted_essay_question_heading_style.font.name = RESTRICTED_RESPONSE_ESSAY_QUESTION_HEADING_STYLES['font'][
            'name']
        restricted_essay_question_heading_style.font.size = RESTRICTED_RESPONSE_ESSAY_QUESTION_HEADING_STYLES['font'][
            'size']
        restricted_essay_question_heading_style.font.color.rgb = RESTRICTED_RESPONSE_ESSAY_QUESTION_HEADING_STYLES[
            'font']['color']


def init_extended_essay_heading(document: Document, style_name: str = 'extended-essay-question-heading'):
    if not is_style_available(document, style_name):
        extended_essay_question_heading_style = document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        extended_essay_question_heading_style.font.name = EXTENDED_RESPONSE_ESSAY_QUESTION_HEADING_STYLES['font'][
            'name']
        extended_essay_question_heading_style.font.size = EXTENDED_RESPONSE_ESSAY_QUESTION_HEADING_STYLES['font'][
            'size']
        extended_essay_question_heading_style.font.color.rgb = EXTENDED_RESPONSE_ESSAY_QUESTION_HEADING_STYLES['font'][
            'color']


def add_cover_page(document: Document) -> None:
    # style definition
    init_cover_page(document)

    # layout definition
    cover_page_run = document.add_paragraph(style='cover-page').add_run()
    cover_page_run.add_picture(COVER_IMAGE_METADATA['filepath'], width=COVER_IMAGE_METADATA['width'],
                               height=COVER_IMAGE_METADATA['height'])
    document.add_page_break()


def init_question_layout(document: Document, subject_area: str) -> None:
    # style definition
    init_topic_heading(document)
    practice_examination_heading_style = document.styles.add_style(
        'practice-examination-heading',
        WD_STYLE_TYPE.PARAGRAPH
    )
    practice_examination_heading_style.paragraph_format.alignment = PRACTICE_EXAMINATION_HEADING_STYLES['paragraph'][
        'alignment']
    practice_examination_heading_style.font.name = PRACTICE_EXAMINATION_HEADING_STYLES['font']['name']
    practice_examination_heading_style.font.size = PRACTICE_EXAMINATION_HEADING_STYLES['font']['size']
    practice_examination_heading_style.font.color.rgb = PRACTICE_EXAMINATION_HEADING_STYLES['font']['color']

    # layout definition
    document.add_paragraph(text=subject_area.upper(), style='topic-heading')
    document.add_paragraph(
        text='Practice Examination',
        style='practice-examination-heading'
    )
    document.add_paragraph()


def init_answer_key_layout(document: Document, subject_area: str) -> None:
    # style definition
    init_topic_heading(document)
    answer_key_heading_style = document.styles.add_style('answer-key-heading', WD_STYLE_TYPE.PARAGRAPH)
    answer_key_heading_style.paragraph_format.alignment = ANSWER_KEY_HEADING_STYLES['paragraph']['alignment']
    answer_key_heading_style.font.name = ANSWER_KEY_HEADING_STYLES['font']['name']
    answer_key_heading_style.font.size = ANSWER_KEY_HEADING_STYLES['font']['size']
    answer_key_heading_style.font.color.rgb = ANSWER_KEY_HEADING_STYLES['font']['color']

    # layout definition
    document.add_paragraph(text=subject_area.upper(), style='topic-heading')
    document.add_paragraph(text='Answer Key', style='answer-key-heading')
    document.add_paragraph()


def add_mcq_question_items(document: Document, question_answers: list[dict]) -> None:
    init_mcq_heading(document)
    document.add_paragraph(
        text='Multiple Choice Questions',
        style='mcq-heading'
    ).add_run(
        text=' (MCQs)',
        style='mcq-heading-abbr'
    )
    for i, mcq in enumerate(question_answers):
        random.shuffle(mcq['alternatives'])
        qa_str = f"{mcq['stem'].strip()}\n\n"
        for j, alt in enumerate(mcq['alternatives']):
            qa_str += f"{chr(ord('A') + j)}. {alt.strip()}"
            if j < len(mcq['alternatives']) - 1:
                qa_str += '\n'
        document.add_paragraph(text=qa_str, style='List Number').paragraph_format.keep_together = True
        if i < len(question_answers) - 1:
            document.add_paragraph()
    document.add_page_break()


def add_mcq_answer_items(document: Document, question_answers: list[dict]) -> None:
    init_mcq_heading(document)
    document.add_paragraph(
        text='Multiple Choice Questions',
        style='mcq-heading'
    ).add_run(
        text=' (MCQs)',
        style='mcq-heading-abbr'
    )
    for mcq in question_answers:
        document.add_paragraph(text=mcq['key'], style='List Number').paragraph_format.keep_together = True
    document.add_page_break()


def add_restricted_essay_question_items(document: Document, question_answers: list[dict]) -> None:
    init_essay_heading(document)
    init_restricted_essay_heading(document)
    document.add_paragraph(text='Essay Questions', style='essay-question-heading').add_run(
        text=' (Restricted Response)',
        style='restricted-essay-question-heading'
    )
    for question_answer in question_answers:
        (document.add_paragraph(text=question_answer['question'], style='List Number')
         .paragraph_format).keep_together = True
    document.add_page_break()


def add_restricted_essay_answer_items(document: Document, question_answers: list[dict]) -> None:
    init_essay_heading(document)
    init_restricted_essay_heading(document)
    document.add_paragraph(text='Essay Questions', style='essay-question-heading').add_run(
        text=' (Restricted Response)',
        style='restricted-essay-question-heading'
    )
    for question_answer in question_answers:
        (document.add_paragraph(text=question_answer['answer'], style='List Number')
         .paragraph_format).keep_together = None
    document.add_page_break()


def add_extended_essay_question_items(document: Document, question_answers: list[dict]) -> None:
    init_essay_heading(document)
    init_extended_essay_heading(document)
    document.add_paragraph(text='Essay Questions', style='essay-question-heading').add_run(
        text=' (Extended Response)',
        style='extended-essay-question-heading'
    )
    for question_answer in question_answers:
        (document.add_paragraph(text=question_answer['question'], style='List Number')
         .paragraph_format).keep_together = True
    document.add_page_break()


def add_extended_essay_answer_items(document: Document, question_answers: list[dict]) -> None:
    init_essay_heading(document)
    init_extended_essay_heading(document)
    document.add_paragraph(text='Essay Questions', style='essay-question-heading').add_run(
        text=' (Extended Response)',
        style='extended-essay-question-heading'
    )
    for question_answer in question_answers:
        (document.add_paragraph(text=question_answer['answer'], style='List Number')
         .paragraph_format).keep_together = None
    document.add_page_break()


def order_documents(cover_page_document: Document, question_layout_document: Document, answer_layout_document: Document,
                    question_documents: list[Document], answer_documents: list[Document]) -> list[Document]:
    ordered_documents = [cover_page_document, question_layout_document]
    for document in question_documents:
        ordered_documents.append(document)
    ordered_documents.append(answer_layout_document)
    for document in answer_documents:
        ordered_documents.append(document)
    return ordered_documents


def merge_documents(documents: list[Document]) -> Composer | None:
    if documents:
        composer = Composer(documents[0])
        for i in range(1, len(documents)):
            if documents[i]:
                composer.append(documents[i])
        return composer


def save(composed_documents: Composer) -> str:
    doc_file_save_path = os.path.join(EXPORT_FILE_SAVE_DIR_PATH,
                                      f'compass-mp-{get_formatted_datetime(datetime.now())}.docx')
    composed_documents.save(doc_file_save_path)
    return doc_file_save_path
