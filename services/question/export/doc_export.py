from docx import Document

from services.question.export.template import ExportService
from utils.question.export.doc_export import (
    add_cover_page,
    init_question_layout,
    init_answer_key_layout,
    add_mcq_question_items,
    add_restricted_essay_question_items,
    add_extended_essay_question_items,
    add_mcq_answer_items,
    add_restricted_essay_answer_items,
    add_extended_essay_answer_items,
    order_documents,
    merge_documents,
    save
)


class DocExportService(ExportService):
    @staticmethod
    def generate_mcq_question_answer_documents(question_documents: list, answer_documents: list,
                                               mcq_question_answers: list[dict]) -> None:
        if mcq_question_answers:
            question_document = Document()
            answer_document = Document()
            add_mcq_question_items(question_document, mcq_question_answers)
            add_mcq_answer_items(answer_document, mcq_question_answers)
            question_documents.append(question_document)
            answer_documents.append(answer_document)

    @staticmethod
    def generate_restricted_essay_question_answer_documents(
            question_documents: list, answer_documents: list,
            restricted_essay_question_answers: list[dict]) -> None:
        if restricted_essay_question_answers:
            question_document = Document()
            answer_document = Document()
            add_restricted_essay_question_items(question_document, restricted_essay_question_answers)
            add_restricted_essay_answer_items(answer_document, restricted_essay_question_answers)
            question_documents.append(question_document)
            answer_documents.append(answer_document)

    @staticmethod
    def generate_extended_essay_question_answer_documents(
            question_documents: list, answer_documents: list,
            extended_essay_question_answers: list[dict]) -> None:
        if extended_essay_question_answers:
            question_document = Document()
            answer_document = Document()
            add_extended_essay_question_items(question_document, extended_essay_question_answers)
            add_extended_essay_answer_items(answer_document, extended_essay_question_answers)
            question_documents.append(question_document)
            answer_documents.append(answer_document)

    @staticmethod
    def generate_export(subject_area: str,
                        mcq_question_answers: list[dict],
                        restricted_essay_question_answers: list[dict],
                        extended_essay_question_answers: list[dict]) -> str:
        question_documents = []
        answer_documents = []
        cover_page_document = Document()
        question_layout_document = Document()
        answer_layout_document = Document()
        add_cover_page(cover_page_document)
        init_question_layout(question_layout_document, subject_area)
        init_answer_key_layout(answer_layout_document, subject_area)
        DocExportService.generate_mcq_question_answer_documents(question_documents, answer_documents,
                                                                mcq_question_answers)
        DocExportService.generate_restricted_essay_question_answer_documents(question_documents, answer_documents,
                                                                             restricted_essay_question_answers)
        DocExportService.generate_extended_essay_question_answer_documents(question_documents, answer_documents,
                                                                           extended_essay_question_answers)
        ordered_documents = order_documents(cover_page_document, question_layout_document, answer_layout_document,
                                            question_documents, answer_documents)
        return save(merge_documents(ordered_documents))
